
from django.shortcuts import render
# from django.contrib.auth.models import User
from . import forms
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# index view  --- > Home Page

def index(request):
    return render( request, 'resumemaker_app/index.html')


# Resume Maker start

#Defining helper functions

def section_head(form, document, s):
    p = document.add_paragraph('')
    p.paragraph_format.space_before = Pt(25)
    r = p.add_run(s)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r.font.name = 'Cambria'
def section_describe(form, document, s):
    para = document.add_paragraph(s)
    para.style = 'List'
def section_subhead(form, document, s):
    p = document.add_paragraph('')
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(s)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r.font.name = 'Cambria'
def add_name(form, document, s):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(s)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 51, 102)
    r.font.name = 'Cambria'
def contact_info(form, document, s):
    p = document.add_paragraph('')
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(s)
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(79, 94, 93)
    r.font.name = 'Cambria'



#Main Function to build resume



def make_resume(request):
    form = forms.resume_form(request.POST)
    if form.is_valid():
        # print(form.cleaned_data['catch'])
        document = Document()

        #Add name --- Title for the resume
        add_name(form, document , form.cleaned_data['full_name'])

        #contact info
        contact_info(form, document, form.cleaned_data['mobile'] + '  |  ' + form.cleaned_data['email'])

        #Summary
        section_head(form, document, 'Summary')
        section_describe(form, document, form.cleaned_data['summary'])

        #Education
        if form.cleaned_data['education_detail_1_1'] != '':
            section_head(form, document, 'Education')
            if form.cleaned_data['education_detail_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['education_detail_1_1'])
                section_describe(form, document,form.cleaned_data['education_detail_1_2'])
                section_describe(form, document, form.cleaned_data['education_detail_1_3'])

            if form.cleaned_data['education_detail_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['education_detail_2_1'])
                section_describe(form, document, form.cleaned_data['education_detail_2_2'])
                section_describe(form, document, form.cleaned_data['education_detail_2_3'])

            if form.cleaned_data['education_detail_3_1'] != '':
                section_subhead(form, document, form.cleaned_data['education_detail_3_1'])
                section_describe(form, document, form.cleaned_data['education_detail_3_2'])
                section_describe(form, document, form.cleaned_data['education_detail_3_3'])

            if form.cleaned_data['education_detail_4_1'] != '':
                section_subhead(form, document, form.cleaned_data['education_detail_4_1'])
                section_describe(form, document, form.cleaned_data['education_detail_4_2'])
                section_describe(form, document,form.cleaned_data['education_detail_4_3'])


        #Work experience

        if form.cleaned_data['experience_1_1'] != '':
            section_head(form , document, 'Work Experience')
            if form.cleaned_data['experience_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['experience_1_1'])
                section_describe(form , document, form.cleaned_data['experience_1_2'])
                section_describe(form , document,form.cleaned_data['experience_1_3'])

            if form.cleaned_data['experience_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['experience_2_1'])
                section_describe(form, document, form.cleaned_data['experience_2_2'])
                section_describe(form, document, form.cleaned_data['experience_2_3'])

            if form.cleaned_data['experience_3_1'] != '':
                section_subhead(form , document, form.cleaned_data['experience_3_1'])
                section_describe(form, document, form.cleaned_data['experience_3_2'])
                section_describe(form, document, form.cleaned_data['experience_3_3'])

            if form.cleaned_data['experience_4_1'] != '':
                section_subhead(form, document, form.cleaned_data['experience_4_1'])
                section_describe(form, document, form.cleaned_data['experience_4_2'])
                section_describe(form, document,form.cleaned_data['experience_4_3'])


        #Certifications

        if form.cleaned_data['certificate_1_1'] != '':
            section_head(form, document, 'Certifications')
            if form.cleaned_data['certificate_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['certificate_1_1'])
                section_describe(form, document, form.cleaned_data['certificate_1_2'])
                section_describe(form, document, form.cleaned_data['certificate_1_3'])

            if form.cleaned_data['certificate_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['certificate_2_1'])
                section_describe(form, document, form.cleaned_data['certificate_2_2'])
                section_describe(form, document,form.cleaned_data['certificate_2_3'])

            if form.cleaned_data['certificate_3_1'] != '':
                section_subhead(form, document, form.cleaned_data['certificate_3_1'])
                section_describe(form, document,form.cleaned_data['certificate_3_2'])
                section_describe(form, document,form.cleaned_data['certificate_3_3'])

            if form.cleaned_data['certificate_4_1'] != '':
                section_subhead(form, document, form.cleaned_data['certificate_4_1'])
                section_describe(form, document,form.cleaned_data['certificate_4_2'])
                section_describe(form, document, form.cleaned_data['certificate_4_3'])


        #Honors and Awards

        if form.cleaned_data['award_1_1'] != '':
            section_head(form, document, 'Honors and Awards')
            if form.cleaned_data['award_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['award_1_1'])
                section_describe(form, document, form.cleaned_data['award_1_2'])

            if form.cleaned_data['award_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['award_2_1'])
                section_describe(form, document, form.cleaned_data['award_2_2'])

            if form.cleaned_data['award_3_1'] != '':
                section_subhead(form, document, form.cleaned_data['award_3_1'])
                section_describe(form, document, form.cleaned_data['award_3_2'])


        #Projects

        if form.cleaned_data['project_1_1'] != '':
            section_head(form, document, 'Projects')
            if form.cleaned_data['project_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['project_1_1'])
                section_describe(form, document, form.cleaned_data['project_1_2'])
                section_describe(form, document, form.cleaned_data['project_1_3'])

            if form.cleaned_data['project_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['project_2_1'])
                section_describe(form, document, form.cleaned_data['project_2_2'])
                section_describe(form, document, form.cleaned_data['project_2_3'])

            if form.cleaned_data['project_3_1'] != '':
                section_subhead(form, document, form.cleaned_data['project_3_1'])
                section_describe(form, document, form.cleaned_data['project_3_2'])
                section_describe(form, document, form.cleaned_data['project_3_3'])

            if form.cleaned_data['project_4_1'] != '':
                section_subhead(form, document, form.cleaned_data['project_4_1'])
                section_describe(form, document, form.cleaned_data['project_4_2'])
                section_describe(form, document, form.cleaned_data['project_4_3'])


        #skills_1_1

        if form.cleaned_data['skills_1_1'] != '':
            section_head(form, document, 'Skills')
            if form.cleaned_data['skills_1_1'] != '':
                section_subhead(form, document, form.cleaned_data['skills_1_1'])
                section_describe(form, document, form.cleaned_data['skills_1_2'])

            if form.cleaned_data['skills_2_1'] != '':
                section_subhead(form, document, form.cleaned_data['skills_2_1'])
                section_describe(form, document, form.cleaned_data['skills_2_2'])


        #Hobbies

        if form.cleaned_data['hobby_1_1'] != '':
            section_head(form, document, 'Hobbies')
            section_describe(form, document, form.cleaned_data['hobby_1_1'])


        response = HttpResponse(document, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=Resume.docx'
        document.save(response)
        return response

    else:
        return render(request, 'resumemaker_app/form.html', {'form' : form})



#Form view function


def resume(request):
    form = forms.resume_form()
    if request.method == 'POST':
        return make_resume(request)

    return render(request, 'resumemaker_app/form.html', {'form' : form})
